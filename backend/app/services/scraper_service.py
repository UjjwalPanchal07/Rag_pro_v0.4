"""
scraper_service.py
==================
Fetches content from any URL and extracts plain text.

Supported sources:
  - Static HTML pages              → requests + BeautifulSoup  (fast)
  - JS-rendered pages (Oracle SPA) → Playwright headless Chrome (slower but complete)
  - PDF (text-based)               → pdfplumber → pypdf fallback
  - PDF (scanned / image-based)    → Claude Haiku Vision OCR
  - DOCX                           → python-docx
  - Excel (.xlsx / .xls)           → openpyxl
  - SharePoint public shared links → same as above (follows redirect)

Installation required for full support:
    pip install playwright pdfplumber python-docx pdf2image
    playwright install chromium
    # For pdf2image on Windows also install Poppler:
    # https://github.com/oschwartz10612/poppler-windows/releases
"""

import base64
import io
import os
import re
import requests
import numpy as np
from bs4 import BeautifulSoup
from dotenv import load_dotenv

from app.models.embedding_model import get_embedding, get_embeddings
from app.models.reranker_model   import rerank

load_dotenv()

# ── Constants ──────────────────────────────────────────────────────────────
PAGE_TIMEOUT_SEC        = 30
CHUNK_SIZE              = 5
TOP_K                   = 10
TOP_CONTEXT             = 1
THRESHOLD               = 0.05   # lowered from 0.10 — catches more relevant chunks
_MAX_WORDS_PER_SENTENCE = 40
_SUB_SPLIT_WORDS        = 30
_MIN_TEXT_LEN           = 200    # below this → suspect JS page or scanned PDF

ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")

_page_cache: dict = {}

_REQUEST_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/120.0.0.0 Safari/537.36"
    ),
    "Accept":          "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
    "Accept-Encoding": "gzip, deflate, br",
    "Connection":      "keep-alive",
}


# ── Content-type detection ─────────────────────────────────────────────────

def _detect_type(url: str, content_type: str) -> str:
    u = url.lower().split("?")[0]
    ct = content_type.lower()
    if u.endswith(".pdf")  or "pdf"              in ct: return "pdf"
    if u.endswith(".docx") or "wordprocessingml" in ct: return "docx"
    if u.endswith(".doc")  or "msword"           in ct: return "docx"
    if u.endswith(".xlsx") or "spreadsheetml"    in ct: return "excel"
    if u.endswith(".xls")  or "ms-excel"         in ct: return "excel"
    return "html"


# ── HTML extraction ────────────────────────────────────────────────────────

def _clean_html(raw) -> str:
    """Strip boilerplate and extract visible text from HTML bytes or string."""
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header",
                     "noscript", "iframe", "form", "aside"]):
        tag.decompose()
    # Also remove elements by common class/id patterns for navigation noise
    for tag in soup.find_all(True, {"class": re.compile(r"nav|sidebar|breadcrumb|toc|menu", re.I)}):
        tag.decompose()
    body = soup.find("body") or soup
    return body.get_text(separator=" ", strip=True)


def _fetch_static(url: str):
    """Fast path — plain HTTP request."""
    r = requests.get(url, headers=_REQUEST_HEADERS, timeout=PAGE_TIMEOUT_SEC, allow_redirects=True)
    r.raise_for_status()
    content_type = r.headers.get("Content-Type", "").lower()
    return r.content, content_type


def _fetch_playwright(url: str) -> str:
    """
    Slow path — headless Chromium for JavaScript-rendered pages (Oracle SPA).
    Falls back gracefully if Playwright is not installed.
    """
    try:
        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            ctx     = browser.new_context(
                user_agent=_REQUEST_HEADERS["User-Agent"],
                locale="en-US",
            )
            page = ctx.new_page()
            page.goto(url, wait_until="domcontentloaded", timeout=30000)

            # Oracle docs load content asynchronously — wait for main content selector
            # Try known Oracle content selectors, fall back to generic wait
            for selector in ["#bookContainer", ".book-content", "article", "main", "#content"]:
                try:
                    page.wait_for_selector(selector, timeout=5000)
                    break
                except Exception:
                    continue

            page.wait_for_timeout(1500)   # extra buffer for lazy-loaded content
            html = page.content()
            browser.close()

        text = _clean_html(html)
        print(f"[scraper] Playwright extracted {len(text):,} chars from {url[:60]}")
        return text

    except ImportError:
        print("[scraper] ⚠️  Playwright not installed.")
        print("[scraper]    Run: pip install playwright && playwright install chromium")
        return ""
    except Exception as e:
        print(f"[scraper] Playwright error for {url[:60]}: {e}")
        return ""


# ── PDF extraction ─────────────────────────────────────────────────────────

def _pdf_vision_ocr(raw: bytes) -> str:
    """Claude Haiku Vision OCR — for scanned/image PDFs."""
    if not ANTHROPIC_API_KEY:
        print("[scraper] ⚠️  ANTHROPIC_API_KEY not set — cannot OCR scanned PDF")
        return ""
    try:
        from pdf2image import convert_from_bytes
    except ImportError:
        print("[scraper] ⚠️  pdf2image not installed — run: pip install pdf2image")
        return ""
    try:
        import anthropic
        images = convert_from_bytes(raw, dpi=200, first_page=1, last_page=5)
        blocks = []
        for img in images:
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            b64 = base64.standard_b64encode(buf.getvalue()).decode()
            blocks.append({"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": b64}})
        blocks.append({"type": "text", "text": "Extract ALL text from these document images exactly as it appears. Include every line, number, date, name, and table cell. Return only the extracted text."})
        client  = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        message = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=2048,
            messages=[{"role": "user", "content": blocks}],
        )
        extracted = message.content[0].text.strip()
        print(f"[scraper] Vision OCR: {len(extracted):,} chars from {len(images)} page(s)")
        return extracted
    except Exception as e:
        print(f"[scraper] Vision OCR failed: {e}")
        return ""


def _extract_pdf(raw: bytes) -> str:
    text = ""
    # 1. pdfplumber (best for text PDFs)
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            parts = [p.extract_text() for p in pdf.pages if p.extract_text()]
        text = " ".join(parts).strip()
    except Exception:
        pass
    # 2. pypdf fallback
    if len(text) < _MIN_TEXT_LEN:
        try:
            from pypdf import PdfReader
            text = " ".join(p.extract_text() or "" for p in PdfReader(io.BytesIO(raw)).pages).strip()
        except Exception:
            pass
    # 3. Vision OCR for scanned PDFs
    if len(text) < _MIN_TEXT_LEN:
        print("[scraper] Scanned PDF detected — using Claude Vision OCR")
        text = _pdf_vision_ocr(raw)
    return text or "Could not extract text from this PDF."


# ── Other file extractors ──────────────────────────────────────────────────

def _extract_docx(raw: bytes) -> str:
    try:
        from docx import Document
        doc   = Document(io.BytesIO(raw))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())
        return " ".join(lines)
    except ImportError:
        return "DOCX extraction unavailable — run: pip install python-docx"
    except Exception as e:
        return f"DOCX extraction error: {e}"


def _extract_excel(raw: bytes) -> str:
    try:
        import openpyxl
        wb    = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
        lines = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(c) for c in row if c is not None and str(c).strip())
                if row_text:
                    lines.append(row_text)
        return " ".join(lines)
    except Exception as e:
        return f"Excel extraction error: {e}"


# ── Main fetch dispatcher ──────────────────────────────────────────────────

def _fetch_text(url: str) -> str:
    """
    Download URL and extract plain text.
    Decision tree:
      1. HTTP GET → detect file type
      2. PDF/DOCX/Excel → dedicated extractor
      3. HTML with enough text → use it directly
      4. HTML with too little text → JS-rendered → Playwright
    """
    try:
        raw, content_type = _fetch_static(url)
    except requests.exceptions.HTTPError as e:
        status = e.response.status_code if e.response else 0
        if status == 403:
            raise Exception("Access denied (403) — this page requires login or blocks automated access.")
        if status == 404:
            raise Exception("Page not found (404) — please check the URL.")
        raise Exception(f"HTTP {status} error fetching the URL.")
    except requests.exceptions.ConnectionError:
        raise Exception("Connection failed — check the URL and your internet connection.")
    except requests.exceptions.Timeout:
        raise Exception("Request timed out — the server is too slow or the URL is unreachable.")

    doc_type = _detect_type(url, content_type)

    # Binary file types — use dedicated extractors
    if doc_type == "pdf":   return _extract_pdf(raw)
    if doc_type == "docx":  return _extract_docx(raw)
    if doc_type == "excel": return _extract_excel(raw)

    # HTML — try static extraction first
    text = _clean_html(raw)
    print(f"[scraper] Static HTML: {len(text):,} chars from {url[:60]}")

    # If too little text → JavaScript SPA (Oracle docs, etc.) → Playwright
    if len(text) < _MIN_TEXT_LEN:
        print(f"[scraper] Only {len(text)} chars from static fetch → trying Playwright")
        js_text = _fetch_playwright(url)
        if len(js_text) > len(text):
            return js_text
        print(f"[scraper] Playwright also returned little text ({len(js_text)} chars) — proceeding with what we have")

    return text


# ── Chunking ───────────────────────────────────────────────────────────────

def _chunk(text: str) -> list[str]:
    parts     = re.split(r"(?<=[.!?])\s+", text)
    sentences = []
    for part in parts:
        part = part.strip()
        if not part or len(part) < 20:
            continue
        words = part.split()
        if len(words) <= _MAX_WORDS_PER_SENTENCE:
            sentences.append(part)
        else:
            for i in range(0, len(words), _SUB_SPLIT_WORDS):
                sub = " ".join(words[i : i + _SUB_SPLIT_WORDS])
                if len(sub) >= 20:
                    sentences.append(sub)
    return [
        " ".join(sentences[i : i + CHUNK_SIZE])
        for i in range(0, len(sentences), CHUNK_SIZE)
    ]


def _load_page(url: str) -> tuple[list[str], np.ndarray]:
    if url in _page_cache:
        return _page_cache[url]
    text   = _fetch_text(url)
    chunks = _chunk(text)
    if not chunks:
        return [], np.empty((0,))
    vecs = get_embeddings(chunks).astype("float32")
    _page_cache[url] = (chunks, vecs)
    return chunks, vecs


# ── Public API ─────────────────────────────────────────────────────────────

def search_from_url(url: str, query: str) -> str:
    try:
        chunks, vecs = _load_page(url)
    except Exception as e:
        return f"Could not fetch the URL: {e}"

    if not chunks:
        return "No usable content found at the provided URL."

    query_vec   = np.array([get_embedding(query)]).astype("float32")
    sims        = (vecs @ query_vec.T).flatten()
    top_k       = min(TOP_K, len(chunks))
    top_indices = np.argsort(sims)[::-1][:top_k]
    candidates  = [chunks[i] for i in top_indices]

    scores   = rerank(query, candidates)
    best_idx = int(np.argmax(scores))

    if scores[best_idx] < THRESHOLD:
        return "No Answer Found at the provided URL."

    ranked_pairs = sorted(zip(scores, candidates), reverse=True)
    top_chunks   = [c for _, c in ranked_pairs[:TOP_CONTEXT]]
    return "\n\n".join(top_chunks)
